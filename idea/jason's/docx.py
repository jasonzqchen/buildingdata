# -*- coding: utf-8 -*-
"""
Created on Wed Jul 25 14:22:47 2018

@author: jason.chen
"""
import docx
import time
import numpy as np
import re
# %% functions.....
def _flatten(l):
    out = []
    for item in l:
        if isinstance(item, (list, tuple)):
            out.extend(_flatten(item))
        else:
            out.append(item)
    return out

# find the table
def _findtable(tables,strings):
    import re
    foundindex = []
    for i in range(len(tables)):
        f_list = _flatten(tables[i])
        list_new = ""
        for item in f_list: 
            list_new = list_new + item
        a = 0
        for string in strings:
            if re.search(string,list_new) is None:
                a += 1
            else:
                a += 0
        if a == 0:
            foundindex.append(i)
    return foundindex



# find position in the table based on input search string
def _findpos(array,str_row,str_col):
    # find col
    from statistics import mode
    (row, col) = aTable.shape
    colf = []
    
    # find col
    colf = [j for string in str_col for i in range(row) for j in range(col) if re.search(string, aTable[i,j]) is not None]
    colf = mode(colf)
    #print(colf)
    
    # find row
    rowf = [i for string in str_row for i in range(row) for j in range(col) if re.search(string, aTable[i,j]) is not None]
    rowf = mode(rowf)
    #print(rowf)
    
    string = aTable[rowf,colf]
    return string

# %% read file

doc = docx.Document(u'test_report.docx')
i = 0
for para in doc.paragraphs:
    if i == 10:
        break
    if para.text != "":
        print(para.text)
        i += 1
        
# %% let's add table to nested list first
start = time.time()     
tables = doc.tables;

list_tables = []
for table in tables:
    list_rows = []
    for row in table.rows:
        list_cells = []
        for cell in row.cells:
            list_cells.append(cell.text)
        list_rows.append(list_cells)
    list_tables.append(list_rows)
print((time.time()-start)*1000) 

# %% main
strings = ['质量','活载','自重|恒载']
index = _findtable(list_tables,strings)
aTable = np.array(list_tables[index[0]])

str_row = ['自重']
str_col = ['YJK','质量']
value = _findpos(aTable, str_row, str_col)
print(value)

# %% test...
# filter by style..
liststyle = []
for paragraph in doc.paragraphs:
    liststyle.append(paragraph.style.name)
    if paragraph.style.name=='Report Level 1':
        print (paragraph.text)
        
# try direct find method...well, it's not faster than putting them in a list...so let's give up this method
start = time.time()
tables = doc.tables;

test = [cell for i in range(len(tables)) for row in tables[i].rows for cell in row.cells if cell.text.find('自重') != -1]
print(test)
print((time.time()-start)*1000) 