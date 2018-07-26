# -*- coding: utf-8 -*-
"""
Created on Wed Jul 25 14:22:47 2018

@author: jason.chen
"""
import docx
import time
# import numpy
# %%
def flatten(l):
    out = []
    for item in l:
        if isinstance(item, (list, tuple)):
            out.extend(flatten(item))
        else:
            out.append(item)
    return out
# %% read file

doc = docx.Document(u'抗震超限专项报告.docx')
i = 0
for para in doc.paragraphs:
    if i == 10:
        break
    if para.text != "":
        print(para.text)
        i += 1
# %% filter by style..
liststyle = []
for paragraph in doc.paragraphs:
    liststyle.append(paragraph.style.name)
    if paragraph.style.name=='Report Level 1':
        print (paragraph.text)
        
# %% try direct find method...well, it's not faster than putting them in a list...so let's give up this method
start = time.time()
tables = doc.tables;

for i in range(len(tables)):
    for row in tables[i].rows:
        for cell in row.cells:
            if cell.text.find('自重') != -1:
                print(i)
print((time.time()-start)*1000) 

# %% let's add table to nested list first
start = time.time()     
tables = doc.tables;

list_tables = []
for table in tables:
    list_rows = []
    for row in table.rows:
        list_cells = []
        for cell in row.cells:
            list_cells.append(cell.text)
        list_rows.append(list_cells)
    list_tables.append(list_rows)
print((time.time()-start)*1000) 
# %% find keywords from a list

# find the table
def findtable(tables,strings):
    import re
    foundindex = []
    for i in range(len(tables)):
        f_list = flatten(tables[i])
        list_new = ""
        for item in f_list: 
            list_new = list_new + item
        a = 0
        for string in strings:
            if re.search(string,list_new) is None:
#            if list_new.find(string) != -1:
                a += 1
            else:
                a += 0
        if a == 0:
            foundindex.append(i)
    return foundindex

# %%
strings = ['质量','活载','自重|恒载']
index = findtable(list_tables,strings)
print(index)

# %%
#for i in range(len(list_tables)):
print(len(index))
for list_row in list_tables[index[0]]:
    
    for item in list_row:
        if item.find('自重') != -1:
            print(list_row)
        if item.find('附加恒载') != -1:
            print(list_row)   
        if item.find('活载') != -1:
            print(list_row)     

# %% let's use numpy....put each table to array...

           

 # %%   
doc.paragraphs    #段落集合
doc.tables        #表格集合
doc.sections      #节  集合 页面格式
doc.styles        #样式集合
doc.inline_shapes #内置图形 等等...

doc.add_paragraph()
doc.save()
    
# %% write to file
import urllib.request
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

urllib.request.urlretrieve("http://placehold.it/350x150", "placeholder.png")
document.add_picture('placeholder.png', width=Inches(1.25))

recordset = [
    {
        "id" : 1,
        "qty": 2,
        "desc": "New item"
    },
    {
        "id" : 2,
        "qty": 2,
        "desc": "New item"
    },
    {
        "id" : 3,
        "qty": 2,
        "desc": "New item"
    },

]

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item["qty"])
    row_cells[1].text = str(item["id"])
    row_cells[2].text = item["desc"]

document.add_page_break()

document.save('demo.docx')