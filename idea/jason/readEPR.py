# -*- coding: utf-8 -*-
"""
Created on Wed Jul 25 14:22:47 2018

@author: jason.chen
@contact: jason.chen@arup.com
@company: ARUP

"""
import docx
import time
import numpy as np
import re

# to solve:
# 1.一开始要筛选一个EPR里有几个项目，分割报告。 难点：doc.paragraphs和tables怎么去关联？？ 要自己重新写一套读word文档的数据库吗？？
# 2.table的定义太乱了，每个EPR都不一样的。。有什么好的方法呢？？
# 3.steel tonnage,从哪里去拿
# 4.柱和墙的尺寸？
# 5.**[IMPORTANT] we need to develop a checking program to check dictionary at the end and to warn user whcih data may not be correct

# output: BuildingDataDict
# it stores all the date got from the file
# units: KN, m

# %% create data dictionary
inputData = ['Project ID','Project Name','Location',\
             'Building Height','Nos of Floors','Building Functions','Structural System',\
             'SW','SDL','LL',\
             'Base_shear_x','Base_moment_x','Base_shear_y','Base_moment_y', \
             'Period:1st Mode','Period:2nd Mode','Period:3rd Mode']
projects = ['project1','project2']

def _createdict(inputdata1, inputdata2):
    value = ''
    datadict = dict()
    for data in inputdata1:
        innerdict = {}
        for key in inputdata2:
            innerdict.update({key:value})
            tempdict = {data:innerdict}
            datadict.update(tempdict)
    return datadict
BuildingDataDict = _createdict(projects,inputData)

# %%
# functions.....
def _flatten(l):
    out = []
    for item in l:
        if isinstance(item, (list, tuple)):
            out.extend(_flatten(item))
        else:
            out.append(item)
    return out

# find the table
def _findtable(tables,strings):
    flag_old = 0
    for i in range(len(tables)):
        f_list = _flatten(tables[i])
        list_new = "".join(f_list)
        flag = 0
        for string in strings:
            if re.search(string,list_new) is None:
                pass
            else:
                flag += 1
                #print('_findtable succesfully found the table...iterating')
        if flag_old < flag:
            foundindex = i
            flag_old = flag
    print('_findtable: found the table index:', foundindex, '\n')
    return foundindex

# find position in the table based on input search string
def _findTablepos(array, str_row, str_col, row_c = []):
    # find col
    from collections import Counter
    (row, col) = array.shape
    colf = []
    
    # find col
    colf = [j for string in str_col for i in range(row) for j in range(col) \
            if re.search(string, array[i,j]) is not None]

    colf = Counter(colf).most_common(1)[0][0]
    #print('colf: ', colf)
    
    # find row
    if row_c == []:
        rowf = [i for string in str_row for i in range(row) for j in range(col) \
                if re.search(string, array[i,j]) is not None]
    else:
        #if row_c to be defined, then only search row in that column
        rowc = [j for string in row_c for j in range(col) \
                if re.search(string, array[0,j]) is not None][0]
        rowf = [i for string in str_row for i in range(row)\
                if re.search(string, array[i,rowc]) is not None]

    rowf = Counter(rowf).most_common(1)[0][0]
    #print('rowf: ', rowf)
    
    string = array[rowf,colf]
    print('_findtpos: located the item in the table:', '[',\
          rowf, ',', colf,']', 'data=', string)
    return string
   
#  let's add table to nested list first
# this is slow...is there any faster method???
def _table2list(tables):
    list_tables = []
    for table in tables:
        list_rows = []
        for row in table.rows:
            list_cells = []
            for cell in row.cells:
                list_cells.append(cell.text)
            list_rows.append(list_cells)
        list_tables.append(list_rows)
    print('_table2list: stored the table into nested list.\n')
    return list_tables

def _outTableData(list_tables, strings, str_row, str_col, row_c=[], datakey=[], \
             datadict={}, project=[]):
    '''
    ***use this functon when dealing with tables***
    list_table: tables from word document in a list
    string: to search for tables
    str_row: keywords to search for rows
    str_col: keywords to search for cols
    row_c: identify which column to locate the rows (optional)
    datakey: key names (if no fill then same as str_row)
    project: project name
    '''
    if datakey == []:
        datakey = str_row
    index = _findtable(list_tables,strings)
    
    # find units
    f_list = _flatten(list_tables[index])
    joinedstring = "".join(f_list)

    # convert to KN
    mulfactor = 1
    if re.search('万吨',joinedstring) is not None:
        mulfactor = 100000
    if re.search('顿|t',joinedstring) is not None:
        mulfactor = 10
    if re.search('MN',joinedstring) is not None:
        mulfactor = 1000
    
    # find table based on keywords
    aTable = np.array(list_tables[index])
    i = 0
    for c in str_col:
        print(c)
        for r in str_row:
            print(r)
            value = _findTablepos(aTable, r, c, row_c)
            
            # apply multiplier for unit
            value = re.sub(',','',value)
            if value.isalpha():
                value = value
            elif type(float(value)) == float:
                value  = '%g' % (float(value) * mulfactor)
                
            if project == []:
                datadict[datakey[i]] = value
            else:
                print(project)
                datadict[project][datakey[i]] = value
            i += 1
            #print(c,r,value)
    print('_outdata: stored the found data into dictionary.\n')
    return datadict

def _parseParaOne(docText, string, splitmark = " "):
    ''' 
    use this function when you want to loop through a text file to find
    the specific 'numbers' or 'words' in one sentence and output that 
    sentence
    DONOT use this function to search keywords that are NOT in ONE sentence!!
    '''
    para = re.search(string,docText)
    parsetext = docText[para.start():para.end()]
    # chinese mark
    parsetext_r = re.sub('[,，。;；]', splitmark, parsetext)
    
    parsedtextlist = parsetext_r.split(splitmark)
    parsedstring = []
    for parsedtext in parsedtextlist:
        temp = re.search(string,parsedtext)
        if temp is not None:
            parsedstring.append(temp.string)
    print('_parsepara found ', len(parsedstring), 'strings: ', parsedstring)
    return parsedstring

def _parsePara(docText, string, keywords):
    para = re.search(string,docText)
    print(para)
    parsetext = docText[para.start():para.end()]
    templist = []
    for keyword in keywords:
        if re.search(keyword,parsetext) is not None:
            keyword = re.sub('[.+*(?)\[\]{}|]','',keyword)
            templist.append(keyword)
    parsedstring = "&".join(templist)
    return parsedstring

# main logic for EPR, to be customised
def _mainEPR(path, datadict):
    doc = docx.Document(path)
    
    docTextlist = [para.text for para in doc.paragraphs]
    docText = " ".join(docTextlist)
    
    list_tables = _table2list(doc.tables)
    
    project = 'project1' ## now use 1 file as example, can loop through multiple
    #datadict = {project:{}}
    
    # project name
    # this best get from folder name while data mining the files
    # 使用文件名， 文件名在 digfiles.py 文档中根据寻找到文件的文件夹名重新命名。
    
    # Location
    string = r'((项目概况)|(工程概括))\s+.+'   
    keywords = ['香港','深圳','上海','北京','广州','天津','武汉','重庆']    
    location = _parsePara(docText, string, keywords)
    datadict[project]['Location'] = location
    
    # building height
    # 结构高度 建筑高度 高度 and 附件有数字 加上m或者米
    string = r'(塔楼)?((楼体高度)|(结构高度)|(建筑高度)|(高度)|高).+\d\d\d\d?.?\d?(m|米)'
    splitpara = _parseParaOne(docText, string)[0]
    height = re.findall("\d+.?\d?",splitpara)[0]
    datadict[project]['Building Height'] = height
    
    # nos. of floors
    # the reason why this is so complicated is I wanted to write a logic that
    # if found multiple xx层 in an array, there's no good way to justify which
    # one is correct, so I used the method estimation nos. of floors comparing 
    # with the found array and return the minimum difference one.
    # the estimated nos. of floors use 4 as storey floor which is a typical
    # floor height for typical buildings
    
    string = r'结?构?高?度?.?\d\d\d?层'
    estfloors = float(height)/4
    splitparas = _parseParaOne(docText, string)
    
    splitpara_n = []
    for splitpara in splitparas:
        splitpara_n.append(float(re.findall('\d+',splitpara)[0]))
    
    adiff = np.array(splitpara_n) - estfloors
    index = np.where(adiff==np.max(adiff))[0][0]
    floors = re.findall(string,splitpara)[index]
    floors = re.findall('\d+',floors)[0]
    datadict[project]['Nos of Floors'] = floors
    
    # building functions..e.g office, comercial..etc
    # 搜索 楼面结构 or 楼板体系 or 楼面体系 办公楼 酒店 公寓
    string = r'((楼面布置)|(楼面结构)|(楼板体系))\s+.+'   
    keywords = ['办公','商业','公寓']    
    functions = _parsePara(docText, string, keywords)
    datadict[project]['Building Functions'] = functions

    # structural system
    # 搜索结构体系， 伸臂， 桁架， 巨柱， 框架， 核心筒， 斜撑
    string = r'(结构体系)\s+.+' 
    keywords = ['矩形钢管混凝土柱','型钢混凝土柱','钢筋混凝土柱'\
            '巨型框架','巨型钢管混凝土柱','巨型型钢混凝土柱','巨型钢筋混凝土柱',\
            '巨型钢?斜支撑(框架)?',\
            '钢筋混凝土梁','钢梁(外框)?',\
            '核心筒',\
            '组合楼板','钢筋混凝土楼?梁?板',\
            '伸臂桁架','腰桁架']
    strsys = _parsePara(docText, string, keywords)
    datadict[project]['Structural System'] = strsys
    
    # steel tonnage <--- how to get??
    # 从给QS的报告里挖？
    # 从模型里挖？
    
    # building period
    # use key words to search tables first, 
    # then use key
    strings = ['振型','周期','YJK','(ETABS)|(Etabs)']
    row_c = ['(周期)|(振型)']
    str_row = ['{:.0f}'.format(x) for x in range(1,4)]
    str_col = [['YJK','周期']]
    datakey = ['Period:1st Mode','Period:2nd Mode','Period:3rd Mode']
    datadict = _outTableData(list_tables,strings,str_row,str_col,row_c,datakey,datadict,project)

    
    # building weight
    strings = ['质量','活载','(自重)|(恒载)','YJK','(ETABS)|(Etabs)']
    row_c = ['项目']
    str_row = ['自重','附加恒载','活载']
    str_col = ['质量']
    datakey = ['SW','SDL','LL']
    datadict = _outTableData(list_tables,strings,str_row,str_col,row_c,datakey,datadict,project)
    
    # wind load and seismic load 
    # 这个还需要仔细研究一下，基本上很多EPR都是不同的格式，可能要写很多个情况。
    strings = ['剪力','(倾覆弯矩)|(倾覆弯矩)','YJK','(ETABS)|(Etabs)']
    row_c = ['荷载工况']
    str_row = [
            ['剪力','(X|x)'],['倾覆弯矩','(X|x)'],['剪力','(Y|y)'],['倾覆弯矩','(Y|y)']
            ]
    str_col = ['YJK']
    datakey = ['Base_shear_x','Base_moment_x','Base_shear_y','Base_moment_y']
    datadict = _outTableData(list_tables,strings,str_row,str_col,row_c,datakey,datadict,project)
    
    return datadict, list_tables

# to be developed
def _checkunit(project, datadict):
    bsx = datadict[project]['Base_shear_x']
    bsy = datadict[project]['Base_shear_y']
    bmx = datadict[project]['Base_moment_x']
    bmy = datadict[project]['Base_moment_y']        
    
# %% running the file
start = time.time()
if __name__ == '__main__':
    path = u'test_report.docx'
    (datadict, list_tables) = _mainEPR(path, BuildingDataDict)

print('run time: ', (time.time()-start)*1000)

# %%

path = u'test_report.docx'
doc = docx.Document(path)
    
docTextlist = [para.text for para in doc.paragraphs]
docText = " ".join(docTextlist)

string = r'(结构体系)\s+.+' 
para = re.search(string,docText)
parsetext = docText[para.start():para.end()]

keywords = ['矩形钢管混凝土柱','型钢混凝土柱','钢筋混凝土柱'\
            '巨型框架','巨型钢管混凝土柱','巨型型钢混凝土柱','巨型钢筋混凝土柱',\
            '巨型钢?斜支撑(框架)?',\
            '钢筋混凝土梁','钢梁(外框)?',\
            '核心筒',\
            '组合楼板','钢筋混凝土楼?梁?板',\
            '伸臂桁架','腰桁架']
functions = []
for keyword in keywords:
    if re.search(keyword,parsetext) is not None:
        keyword = re.sub('[.+*(?)\[\]{}|]','',keyword)
        functions.append(keyword)
"&".join(functions)
